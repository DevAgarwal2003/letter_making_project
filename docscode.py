import streamlit as st
from mailmerge import MailMerge
import pandas as pd
from io import BytesIO
from zipfile import ZipFile
from docx import Document

# Function to transform column names
def transform_column_names(col):
    col = col.replace('\n', '_') if '\n' in col and not col.startswith('\n') else col
    col = col.replace('\n', 'M__', 1) if col.startswith('\n') else col
    col = ''.join(char for char in col if char not in '()')
    col = col.replace('/', '').replace('.', '')
    return col

# Function to replace text in Word document while preserving bold formatting
def replace_text_preserve_formatting(doc, old_text, new_text):
    for para in doc.paragraphs:
        for run in para.runs:
            if old_text in run.text:
                is_bold = run.bold  # Check if the original text is bold
                run.text = run.text.replace(old_text, new_text)
                run.bold = is_bold  # Apply the same bold formatting
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old_text in run.text:
                            is_bold = run.bold
                            run.text = run.text.replace(old_text, new_text)
                            run.bold = is_bold
    return doc

def clean_value(value):
    """Convert NaN to an empty string and strip spaces."""
    if pd.isna(value):  # Check if value is NaN
        return ""
    return str(value).strip()

# Function to perform mail merge
def perform_mail_merge(word_file, excel_file):
    data = pd.read_excel(excel_file)
    data.columns = data.columns.str.replace(' ', '_')
    data.columns = [transform_column_names(col) for col in data.columns]

    # Identify datetime columns and format them to DD-MM-YYYY
    for col in data.columns:
        if pd.api.types.is_datetime64_any_dtype(data[col]):
            data[col] = pd.to_datetime(data[col]).dt.strftime('%d-%m-%Y')

    # Preprocess text columns
    for col in data.columns:
        if data[col].dtype == 'object':
            data[col] = data[col].str.strip()
            data[col] = data[col].str.replace(r'\s+', ' ', regex=True)
            data[col] = data[col].str.replace(r'\n|\r', ' ', regex=True)

    output_files = []

    for index, row in data.iterrows():
        merge_fields = {field: str(row[field]) for field in row.index}
        
        # Determine 'option' value dynamically
        borrower_name = clean_value(row.get("Borrowers_Name", ""))
        borrower_2_name = clean_value(row.get("Borrower_2_Name", ""))
        borrower_3_name = clean_value(row.get("Borrower_3_Name", ""))
        if borrower_name and borrower_2_name and borrower_3_name:  
            option_value = "AND ORS"  # All three exist
        elif borrower_name and borrower_2_name and not borrower_3_name:  
            option_value = "AND ANR"  # Only Borrower 1 and 2 exist
        elif borrower_name and not borrower_2_name and not borrower_3_name:  
            option_value = ""  # Only Borrower 1 exists
        else:  
            option_value = "AND ANR"  # Any other case (e.g., Borrower 1 and Borrower 3 exist, or only Borrower 2 exists)

        with MailMerge(word_file) as document:
            document.merge(**merge_fields)
            output_stream = BytesIO()
            document.write(output_stream)
            output_stream.seek(0)

        # Load the merged document and replace <<option>> manually while keeping bold
        doc = Document(output_stream)
        doc = replace_text_preserve_formatting(doc, "<<option>>", option_value)

        # Save the final document
        final_output_stream = BytesIO()
        doc.save(final_output_stream)
        final_output_stream.seek(0)

        output_files.append((f"Document_{index + 1}.docx", final_output_stream))

    return output_files

# Streamlit UI
st.markdown("""
    <style>
    body {
        background-color: black;
    }
    .title-style {
        font-size: 48px;
        color: white;
        background-color: #D2B48C;  
        padding: 20px;
        text-align: center;
        border-radius: 10px;
        width: 100%;
        margin: 0 auto;
    }
    </style>
    <div class="title-style">
        Letters Drafting App
    </div>
    """, unsafe_allow_html=True)

# Upload Word template file
word_file = st.file_uploader("📤 Upload Word Template (.docx)", type=["docx"])

# Upload Excel file
excel_file = st.file_uploader("📤 Upload Excel File (.xlsx)", type=["xlsx"])

if word_file and excel_file:
    if st.button("Generate Documents"):
        with st.spinner("Generating documents..."):
            try:
                output_files = perform_mail_merge(word_file, excel_file)
                zip_buffer = BytesIO()
                with ZipFile(zip_buffer, "w") as zip_file:
                    for filename, file_stream in output_files:
                        zip_file.writestr(filename, file_stream.getvalue())
                zip_buffer.seek(0)

                st.success("Documents generated successfully!")
                st.download_button(
                    label="Download All Documents",
                    data=zip_buffer,
                    file_name="Generated_Documents.zip",
                    mime="application/zip",
                )
            except Exception as e:
                st.error(f"An error occurred: {e}")
